"""Render a Fiberarticle document to .docx in journal-style templates.

Templates:
- generic: single column, Calibri, numbered headings, numeric [n] citations.
- ieee: Times New Roman 10pt body, two-column layout, IEEE-style headings and
  numeric reference list.
- apa: Times New Roman 12pt, double spacing, APA 7 manuscript conventions,
  author-date in-text citations (converted from [n] markers) and a hanging
  indent reference list.

The document JSON remains the source of truth; the .docx is only a render.
"""

import io
import re

from docx import Document as DocxDocument
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

_CITE_RE = re.compile(r"\[(\d+(?:\s*,\s*\d+)*)\]")
_PLACEHOLDER_RE = re.compile(r"^\[.*placeholder.*\]$", re.IGNORECASE)

# Block-level Markdown: list items and blockquotes.
_LIST_BULLET_RE = re.compile(r"^[-*+]\s+(.+)$")
_LIST_NUMBER_RE = re.compile(r"^\d+[.)]\s+(.+)$")
_BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")

# Inline Markdown marks, tried in priority order (bold before italic so
# "**x**" is not swallowed by the single-marker italic alternatives).
_INLINE_TOKEN_RE = re.compile(
    r"\*\*(?P<bold>.+?)\*\*"
    r"|__(?P<bold2>.+?)__"
    r"|\*(?P<italic>[^*]+?)\*"
    r"|_(?P<italic2>[^_]+?)_"
    r"|~~(?P<strike>.+?)~~"
    r"|`(?P<code>[^`]+?)`"
)


def _surname(author: str) -> str:
    parts = author.strip().split()
    return parts[-1] if parts else author


def _set_two_columns(section) -> None:
    cols = section._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "360")


def _parse_inline(text: str) -> list[tuple[str, frozenset]]:
    """Tokenize inline Markdown marks into (text, marks) segments.

    Recurses into matched spans so nesting like ``**bold *italic***``
    accumulates marks on the inner segment. Unmatched/unbalanced markers
    (no closing pair found) are left as literal text rather than raising.
    """
    segments: list[tuple[str, frozenset]] = []
    pos = 0
    for m in _INLINE_TOKEN_RE.finditer(text):
        if m.start() < pos:
            continue
        if m.start() > pos:
            segments.append((text[pos : m.start()], frozenset()))
        if m.group("bold") is not None:
            inner, mark = m.group("bold"), "bold"
        elif m.group("bold2") is not None:
            inner, mark = m.group("bold2"), "bold"
        elif m.group("italic") is not None:
            inner, mark = m.group("italic"), "italic"
        elif m.group("italic2") is not None:
            inner, mark = m.group("italic2"), "italic"
        elif m.group("strike") is not None:
            inner, mark = m.group("strike"), "strike"
        else:
            inner, mark = m.group("code"), "code"
        for seg_text, seg_marks in _parse_inline(inner):
            segments.append((seg_text, seg_marks | {mark}))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], frozenset()))
    return segments


def _write_runs(
    p, text: str, style: dict, base_marks: frozenset = frozenset()
) -> None:
    """Write text into paragraph p as runs, applying inline Markdown marks."""
    for seg_text, seg_marks in _parse_inline(text):
        if not seg_text:
            continue
        marks = base_marks | seg_marks
        run = p.add_run(seg_text)
        run.font.name = "Consolas" if "code" in marks else style["font"]
        run.font.size = style["body_size"]
        if "bold" in marks:
            run.bold = True
        if "italic" in marks:
            run.italic = True
        if "strike" in marks:
            run.font.strike = True


def _add_body_paragraphs(
    doc, text: str, style: dict, intext: dict[str, str] | None
):
    try:
        doc.styles["Intense Quote"]
        quote_style_exists = True
    except KeyError:
        quote_style_exists = False

    for raw in text.split("\n"):
        line = raw.strip()
        if not line:
            continue
        if intext:
            line = _CITE_RE.sub(lambda m: intext.get(m.group(0), m.group(0)), line)

        list_style = None
        base_marks: frozenset = frozenset()

        m = _LIST_BULLET_RE.match(line)
        if m:
            list_style, line = "List Bullet", m.group(1)
        else:
            m = _LIST_NUMBER_RE.match(line)
            if m:
                list_style, line = "List Number", m.group(1)
            else:
                m = _BLOCKQUOTE_RE.match(line)
                if m:
                    line = m.group(1)
                    if quote_style_exists:
                        list_style = "Intense Quote"
                    else:
                        base_marks = frozenset({"italic"})

        p = doc.add_paragraph()
        if list_style:
            try:
                p.style = doc.styles[list_style]
            except KeyError:
                pass

        is_placeholder = list_style is None and bool(_PLACEHOLDER_RE.match(line))
        if is_placeholder:
            base_marks = base_marks | {"italic"}

        _write_runs(p, line, style, base_marks)

        if is_placeholder:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif list_style is None:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = style["indent"]
        if style["double_space"]:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_after = style["space_after"]


_STYLES = {
    "generic": {
        "font": "Calibri",
        "body_size": Pt(11),
        "title_size": Pt(18),
        "heading_size": Pt(13),
        "double_space": False,
        "indent": None,
        "space_after": Pt(8),
        "numbered_headings": True,
        "two_column": False,
    },
    "ieee": {
        "font": "Times New Roman",
        "body_size": Pt(10),
        "title_size": Pt(24),
        "heading_size": Pt(10),
        "double_space": False,
        "indent": Inches(0.17),
        "space_after": Pt(4),
        "numbered_headings": True,
        "two_column": True,
    },
    "apa": {
        "font": "Times New Roman",
        "body_size": Pt(12),
        "title_size": Pt(12),
        "heading_size": Pt(12),
        "double_space": True,
        "indent": Inches(0.5),
        "space_after": Pt(0),
        "numbered_headings": False,
        "two_column": False,
    },
}

_ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]

# Templates whose real output is LaTeX reuse the closest .docx layout.
_LAYOUT_FOR_TEMPLATE = {
    "generic": "generic",
    "ieee": "ieee",
    "apa": "apa",
    "acm": "ieee",
    "elsevier": "generic",
    "springer": "generic",
    "neurips": "generic",
}


def render_docx(
    document: dict,
    papers: list[dict],
    references: list[str] | None = None,
    intext: dict[str, str] | None = None,
    numeric: bool = True,
) -> bytes:
    """Render the document JSON to .docx.

    references: pre-rendered reference entries (CSL engine), one per paper,
    in paper order. intext: replacement map for bracketed citation markers,
    used by author-date styles; numeric styles keep [n] markers.
    """
    template = document.get("template") or "generic"
    layout = _LAYOUT_FOR_TEMPLATE.get(template, "generic")
    style = _STYLES.get(layout, _STYLES["generic"])
    apa = layout == "apa"

    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1 if not style["two_column"] else 0.75)
    section.right_margin = Inches(1 if not style["two_column"] else 0.75)

    normal = doc.styles["Normal"]
    normal.font.name = style["font"]
    normal.font.size = style["body_size"]
    # East Asian font fallback so Word does not substitute.
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), style["font"])

    # Title block (single column, before any column break)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(document["title"])
    title_run.font.name = style["font"]
    title_run.font.size = style["title_size"]
    title_run.bold = True
    if style["double_space"]:
        title_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    authors = document.get("authors") or []
    if authors:
        author_p = doc.add_paragraph()
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_p.add_run(", ".join(authors))
        author_run.font.name = style["font"]
        author_run.font.size = style["body_size"]

    if style["two_column"]:
        # IEEE-style: title spans the page, body flows in two columns.
        body_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
        body_section.left_margin = Inches(0.75)
        body_section.right_margin = Inches(0.75)
        _set_two_columns(body_section)

    heading_counter = 0
    for sec in document.get("sections") or []:
        heading = sec.get("heading") or "Section"
        content = sec.get("content") or ""

        is_abstract = heading.strip().lower() == "abstract"
        h = doc.add_paragraph()
        if style["numbered_headings"] and not is_abstract:
            heading_counter += 1
            if layout == "ieee":
                label = f"{_ROMAN[min(heading_counter - 1, len(_ROMAN) - 1)]}. {heading.upper()}"
            else:
                label = f"{heading_counter}. {heading}"
        else:
            label = heading
        h_run = h.add_run(label)
        h_run.font.name = style["font"]
        h_run.font.size = style["heading_size"]
        h_run.bold = True
        if apa:
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_abstract else WD_ALIGN_PARAGRAPH.LEFT
            h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)

        _add_body_paragraphs(doc, content, style, intext if not numeric else None)

    # References
    ref_h = doc.add_paragraph()
    ref_label = "References"
    if layout == "ieee":
        ref_label = "REFERENCES"
    ref_run = ref_h.add_run(ref_label)
    ref_run.font.name = style["font"]
    ref_run.font.size = style["heading_size"]
    ref_run.bold = True
    if apa:
        ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    ref_h.paragraph_format.space_before = Pt(10)

    if references is None:
        # Minimal fallback when no engine-rendered entries were provided.
        references = [
            f"{', '.join((p.get('authors') or [])[:6])} "
            f"({p.get('year') or 'n.d.'}). {p['title']}.".strip()
            for p in papers
        ]

    if numeric:
        entries = [f"[{i}] {text}" for i, text in enumerate(references, 1)]
    else:
        # Author-date reference lists are alphabetical.
        entries = sorted(references, key=str.lower)

    for text in entries:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = style["font"]
        run.font.size = style["body_size"]
        if not numeric:
            # Hanging indent, the convention for author-date lists.
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
        if apa:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_after = Pt(4)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
